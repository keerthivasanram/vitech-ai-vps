"""Structure-preserving document extraction: PDF/DOCX/XLSX/JSON/TXT.

Unlike a flat text dump, this returns a list of Blocks, each carrying:
  * its page number (PDF only — DOCX/XLSX have no fixed pages)
  * whether it is prose (`text`) or a `table`

Tables matter: a Vitech offer's technical specification and price schedule are
tables, and the sizing numbers live in them. pdfplumber extracts them as real
rows (rendered here as pipe-delimited lines) and — crucially — the surrounding
prose is extracted with the table regions removed, so a table's cells are not
also smeared into the body text as garbled fragments.
"""
from __future__ import annotations

import json
from dataclasses import dataclass, field
from pathlib import Path

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".json", ".txt", ".md"}

TABLE_MARKER = "[TABLE]"


@dataclass
class Block:
    text: str
    page: int | None = None          # 1-indexed PDF page; None where not paged
    kind: str = "text"               # "text" | "table"
    meta: dict = field(default_factory=dict)


# --- table rendering -------------------------------------------------------

def _render_table(rows: list[list]) -> str:
    """Pipe-delimited rows, blank cells kept as empty so columns stay aligned
    positionally. Newlines inside a cell are flattened."""
    lines = []
    for row in rows:
        cells = ["" if c is None else str(c).strip().replace("\n", " ") for c in row]
        if any(cells):
            lines.append(" | ".join(cells))
    return "\n".join(lines)


# --- PDF -------------------------------------------------------------------

def load_pdf_blocks(path: Path) -> list[Block]:
    import pdfplumber

    blocks: list[Block] = []
    with pdfplumber.open(str(path)) as pdf:
        for page_no, page in enumerate(pdf.pages, start=1):
            tables = page.find_tables()
            bboxes = [t.bbox for t in tables]

            def _outside_tables(obj, _bboxes=bboxes):
                top, bottom, x0, x1 = obj.get("top"), obj.get("bottom"), obj.get("x0"), obj.get("x1")
                if top is None:
                    return True
                for (bx0, btop, bx1, bbottom) in _bboxes:
                    if top >= btop and bottom <= bbottom and x0 >= bx0 and x1 <= bx1:
                        return False
                return True

            prose = (page.filter(_outside_tables).extract_text() or "").strip() if bboxes \
                else (page.extract_text() or "").strip()
            if prose:
                blocks.append(Block(text=prose, page=page_no, kind="text"))

            for t in tables:
                rendered = _render_table(t.extract())
                if rendered:
                    blocks.append(Block(text=f"{TABLE_MARKER}\n{rendered}",
                                        page=page_no, kind="table"))
    return blocks


# --- DOCX ------------------------------------------------------------------

def load_docx_blocks(path: Path) -> list[Block]:
    import docx

    doc = docx.Document(str(path))
    blocks: list[Block] = []
    # Paragraphs and tables in document order (python-docx exposes them
    # separately, so walk the body XML to keep their true sequence).
    body = doc.element.body
    para_map = {p._p: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}
    for child in body.iterchildren():
        if child in para_map:
            text = para_map[child].text.strip()
            if text:
                blocks.append(Block(text=text, kind="text"))
        elif child in table_map:
            rows = [[c.text for c in row.cells] for row in table_map[child].rows]
            rendered = _render_table(rows)
            if rendered:
                blocks.append(Block(text=f"{TABLE_MARKER}\n{rendered}", kind="table"))
    return blocks


# --- XLSX ------------------------------------------------------------------

def load_xlsx_blocks(path: Path) -> list[Block]:
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    blocks: list[Block] = []
    for sheet in wb.worksheets:
        rows = [list(r) for r in sheet.iter_rows(values_only=True)]
        rendered = _render_table(rows)
        if rendered:
            blocks.append(Block(text=f"{TABLE_MARKER} Sheet: {sheet.title}\n{rendered}",
                                kind="table", meta={"sheet": sheet.title}))
    return blocks


# --- TXT / MD / JSON -------------------------------------------------------

def load_txt_blocks(path: Path) -> list[Block]:
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    return [Block(text=text, kind="text")] if text else []


def _flatten_json(value, prefix: str = "") -> list[str]:
    lines: list[str] = []
    if isinstance(value, dict):
        for k, v in value.items():
            label = f"{prefix}.{k}" if prefix else str(k)
            if isinstance(v, (dict, list)):
                lines.extend(_flatten_json(v, label))
            elif v not in (None, "", []):
                lines.append(f"{label}: {v}")
    elif isinstance(value, list):
        for i, item in enumerate(value):
            lines.extend(_flatten_json(item, f"{prefix}[{i}]"))
    return lines


def load_json_blocks(path: Path) -> list[Block]:
    data = json.loads(path.read_text(encoding="utf-8"))
    text = "\n".join(_flatten_json(data)).strip()
    return [Block(text=text, kind="text")] if text else []


_LOADERS = {
    ".pdf": load_pdf_blocks,
    ".docx": load_docx_blocks,
    ".xlsx": load_xlsx_blocks,
    ".xls": load_xlsx_blocks,
    ".json": load_json_blocks,
    ".txt": load_txt_blocks,
    ".md": load_txt_blocks,
}


def load_blocks(path: Path) -> list[Block]:
    """Dispatch on extension; return the document's structure-preserving blocks."""
    ext = path.suffix.lower()
    loader = _LOADERS.get(ext)
    if loader is None:
        raise ValueError(f"{path}: unsupported extension {ext!r} "
                         f"(supported: {sorted(SUPPORTED_EXTENSIONS)})")
    return loader(path)


def load_text(path: Path) -> str:
    """Whole-document plain text (blocks joined) — used for metadata scanning."""
    return "\n".join(b.text for b in load_blocks(path))


def iter_source_files(source: Path):
    """Yield every supported file under `source` (a file or a directory)."""
    if source.is_dir():
        for path in sorted(source.rglob("*")):
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                yield path
    elif source.suffix.lower() in SUPPORTED_EXTENSIONS:
        yield source
    else:
        raise ValueError(f"{source}: unsupported extension {source.suffix!r}")
